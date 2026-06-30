"""
document_profiler.py

Production V1 - Colab Ready

Purpose
-------
Analyze document characteristics before deep extraction.

This module answers:
- What type of document is this?
- Is it text-based PDF, scanned PDF, or hybrid PDF?
- Does it need OCR?
- Which pages likely need OCR?
- Does it contain tables/images/drawings?
- What processing strategy should downstream pipeline use?

Input
-----
document_path

Output
------
document_profile.json

Typical usage
-------------
from document_ai.ingestion.document_profiler import DocumentProfiler

profiler = DocumentProfiler()
profile = profiler.process("sample.pdf")

print(profile.to_json(ensure_ascii=False, indent=2))
"""

from __future__ import annotations

import json
import statistics
from pathlib import Path
from dataclasses import dataclass, asdict, field
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional, Union


# ============================================================
# Local import
# ============================================================

try:
    from document_ai.ingestion.document_loader import (
        DocumentLoader,
        DocumentLoaderConfig,
        LoadedDocument,
        MissingDependencyError,
        DocumentLoaderError,
    )
except Exception:
    DocumentLoader = None
    DocumentLoaderConfig = None
    LoadedDocument = None

    class MissingDependencyError(Exception):
        pass

    class DocumentLoaderError(Exception):
        pass


# ============================================================
# Optional dependencies
# ============================================================

try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

try:
    from docx import Document as DocxDocument
except Exception:
    DocxDocument = None

try:
    from PIL import Image
except Exception:
    Image = None


# ============================================================
# Config
# ============================================================

@dataclass
class DocumentProfilerConfig:
    """
    Configuration for DocumentProfiler.
    """

    sample_pages: int = 10

    # PDF text threshold
    min_text_chars_for_text_layer: int = 30
    min_words_for_text_page: int = 5

    # OCR detection
    ocr_text_char_threshold: int = 20
    ocr_image_area_ratio_threshold: float = 0.35

    # Table heuristic
    table_min_aligned_lines: int = 3
    table_min_numeric_lines: int = 3

    # Complexity thresholds
    high_complexity_table_ratio: float = 0.35
    high_complexity_image_ratio: float = 0.50
    high_complexity_avg_blocks: int = 80

    # Whether to include per-page profiles in final output
    include_page_profiles: bool = True


# ============================================================
# Schemas
# ============================================================

@dataclass
class PageProfile:
    """
    Per-page profile.
    """

    page_number: int
    width: float
    height: float
    rotation: int

    text_char_count: int = 0
    word_count: int = 0
    text_block_count: int = 0
    image_count: int = 0
    drawing_count: int = 0

    has_text: bool = False
    has_image: bool = False
    has_drawing: bool = False
    likely_scanned: bool = False
    need_ocr: bool = False

    table_candidate_count: int = 0
    has_table_candidate: bool = False

    text_density: float = 0.0
    image_area_ratio: float = 0.0

    page_type: str = "unknown"
    layout_complexity: str = "low"

    warnings: List[str] = field(default_factory=list)

    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)


@dataclass
class DocumentProfile:
    """
    JSON-safe document profile.
    """

    document_id: str
    source_path: str
    file_name: str
    file_extension: str
    document_type: str
    mime_type: str
    file_size_bytes: int
    sha256: str

    page_count: Optional[int] = None
    is_encrypted: bool = False

    has_text_layer: Optional[bool] = None
    has_image_layer: Optional[bool] = None
    need_ocr: Optional[bool] = None

    pdf_type: Optional[str] = None
    language_hint: Optional[str] = None
    processing_strategy: str = "unknown"
    complexity_level: str = "unknown"

    ocr_page_numbers: List[int] = field(default_factory=list)
    table_candidate_pages: List[int] = field(default_factory=list)
    image_pages: List[int] = field(default_factory=list)

    stats: Dict[str, Any] = field(default_factory=dict)
    metadata: Dict[str, Any] = field(default_factory=dict)
    pages: List[Dict[str, Any]] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)

    profiled_at_utc: str = field(
        default_factory=lambda: datetime.now(timezone.utc).isoformat()
    )

    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)

    def to_json(self, ensure_ascii: bool = False, indent: int = 2) -> str:
        return json.dumps(self.to_dict(), ensure_ascii=ensure_ascii, indent=indent)


# ============================================================
# Document Profiler
# ============================================================

class DocumentProfiler:
    """
    Analyze document before extraction.

    This class should not perform full document extraction.
    It only profiles document characteristics and recommends
    a downstream processing strategy.
    """

    def __init__(
        self,
        config: Optional[DocumentProfilerConfig] = None,
        loader: Optional[Any] = None,
    ):
        self.config = config or DocumentProfilerConfig()

        if loader is not None:
            self.loader = loader
        else:
            if DocumentLoader is None:
                raise MissingDependencyError(
                    "DocumentLoader is not available. "
                    "Make sure document_ai/ingestion/document_loader.py exists."
                )

            self.loader = DocumentLoader(
                DocumentLoaderConfig(
                    allow_encrypted_pdf=False,
                    sample_pages_for_pdf_profile=self.config.sample_pages,
                )
            )

    # --------------------------------------------------------
    # Public API
    # --------------------------------------------------------

    def process(
        self,
        document_path: Union[str, Path],
        document_id: Optional[str] = None,
        profile_all_pages: bool = False,
    ) -> DocumentProfile:
        """
        Profile a document.

        Parameters
        ----------
        document_path:
            Input document path.

        document_id:
            Optional custom document id.

        profile_all_pages:
            If True, profile every PDF page.
            If False, profile selected sample pages.

        Returns
        -------
        DocumentProfile
        """

        loaded_doc = self.loader.load(document_path, document_id=document_id)

        base = self._create_base_profile(loaded_doc)

        if loaded_doc.document_type == "pdf":
            return self._profile_pdf(
                document_path=document_path,
                base=base,
                loaded_doc=loaded_doc,
                profile_all_pages=profile_all_pages,
            )

        if loaded_doc.document_type == "docx":
            return self._profile_docx(
                document_path=document_path,
                base=base,
                loaded_doc=loaded_doc,
            )

        if loaded_doc.document_type == "image":
            return self._profile_image(
                document_path=document_path,
                base=base,
                loaded_doc=loaded_doc,
            )

        if loaded_doc.document_type == "text":
            return self._profile_text(
                document_path=document_path,
                base=base,
                loaded_doc=loaded_doc,
            )

        base.warnings.append(
            f"No profiler implemented for type: {loaded_doc.document_type}"
        )
        return base

    # --------------------------------------------------------
    # Base profile
    # --------------------------------------------------------

    def _create_base_profile(self, loaded_doc: Any) -> DocumentProfile:
        return DocumentProfile(
            document_id=loaded_doc.document_id,
            source_path=loaded_doc.source_path,
            file_name=loaded_doc.file_name,
            file_extension=loaded_doc.file_extension,
            document_type=loaded_doc.document_type,
            mime_type=loaded_doc.mime_type,
            file_size_bytes=loaded_doc.file_size_bytes,
            sha256=loaded_doc.sha256,
            page_count=loaded_doc.page_count,
            is_encrypted=loaded_doc.is_encrypted,
            has_text_layer=loaded_doc.has_text_layer,
            has_image_layer=loaded_doc.has_image_layer,
            need_ocr=loaded_doc.need_ocr,
            metadata={
                "loader_metadata": loaded_doc.metadata,
            },
            warnings=list(loaded_doc.warnings or []),
        )

    # --------------------------------------------------------
    # PDF profiling
    # --------------------------------------------------------

    def _profile_pdf(
        self,
        document_path: Union[str, Path],
        base: DocumentProfile,
        loaded_doc: Any,
        profile_all_pages: bool = False,
    ) -> DocumentProfile:
        if fitz is None:
            raise MissingDependencyError(
                "PyMuPDF is not installed. Run: pip install pymupdf"
            )

        path = Path(document_path)

        with fitz.open(str(path)) as doc:
            if doc.is_encrypted:
                raise DocumentLoaderError(
                    "PDF is encrypted. Cannot profile without password."
                )

            page_indices = self._select_page_indices(
                page_count=doc.page_count,
                profile_all_pages=profile_all_pages,
            )

            page_profiles: List[PageProfile] = []

            for page_index in page_indices:
                try:
                    page = doc.load_page(page_index)
                    page_profile = self._profile_pdf_page(page, page_index)
                    page_profiles.append(page_profile)
                except Exception as exc:
                    base.warnings.append(
                        f"Failed to profile page {page_index + 1}: {exc}"
                    )

            self._aggregate_pdf_profile(base, page_profiles, doc.page_count)

            if self.config.include_page_profiles:
                base.pages = [p.to_dict() for p in page_profiles]

            base.metadata.update(
                {
                    "pdf_metadata": doc.metadata or {},
                    "profiled_page_count": len(page_profiles),
                    "profiled_page_numbers": [
                        p.page_number for p in page_profiles
                    ],
                    "profile_all_pages": profile_all_pages,
                }
            )

        return base

    def _profile_pdf_page(self, page: Any, page_index: int) -> PageProfile:
        rect = page.rect
        page_area = max(float(rect.width * rect.height), 1.0)

        warnings: List[str] = []

        text = ""
        text_dict: Dict[str, Any] = {}
        drawings = []
        images = []

        try:
            text = page.get_text("text") or ""
        except Exception as exc:
            warnings.append(f"Cannot extract text: {exc}")

        try:
            text_dict = page.get_text("dict") or {}
        except Exception as exc:
            warnings.append(f"Cannot extract text dict: {exc}")

        try:
            images = page.get_images(full=True) or []
        except Exception as exc:
            warnings.append(f"Cannot inspect images: {exc}")

        try:
            drawings = page.get_drawings() or []
        except Exception as exc:
            warnings.append(f"Cannot inspect drawings: {exc}")

        text_char_count = len(text.strip())
        words = text.split()
        word_count = len(words)

        text_block_count = self._count_text_blocks(text_dict)
        image_count = len(images)
        drawing_count = len(drawings)

        image_area_ratio = self._estimate_image_area_ratio(page, page_area)

        table_candidate_count = self._estimate_table_candidates(
            text=text,
            drawings=drawings,
        )

        has_text = (
            text_char_count >= self.config.min_text_chars_for_text_layer
            and word_count >= self.config.min_words_for_text_page
        )
        has_image = image_count > 0
        has_drawing = drawing_count > 0
        has_table_candidate = table_candidate_count > 0

        likely_scanned = (
            not has_text
            and has_image
            and image_area_ratio >= self.config.ocr_image_area_ratio_threshold
        )

        need_ocr = (
            text_char_count < self.config.ocr_text_char_threshold
            and (
                has_image
                or image_area_ratio >= self.config.ocr_image_area_ratio_threshold
            )
        )

        text_density = text_char_count / page_area

        page_type = self._classify_page_type(
            has_text=has_text,
            has_image=has_image,
            has_table_candidate=has_table_candidate,
            likely_scanned=likely_scanned,
        )

        layout_complexity = self._classify_page_complexity(
            text_block_count=text_block_count,
            image_count=image_count,
            drawing_count=drawing_count,
            table_candidate_count=table_candidate_count,
        )

        return PageProfile(
            page_number=page_index + 1,
            width=float(rect.width),
            height=float(rect.height),
            rotation=int(page.rotation),
            text_char_count=text_char_count,
            word_count=word_count,
            text_block_count=text_block_count,
            image_count=image_count,
            drawing_count=drawing_count,
            has_text=has_text,
            has_image=has_image,
            has_drawing=has_drawing,
            likely_scanned=likely_scanned,
            need_ocr=need_ocr,
            table_candidate_count=table_candidate_count,
            has_table_candidate=has_table_candidate,
            text_density=text_density,
            image_area_ratio=image_area_ratio,
            page_type=page_type,
            layout_complexity=layout_complexity,
            warnings=warnings,
        )

    def _aggregate_pdf_profile(
        self,
        base: DocumentProfile,
        page_profiles: List[PageProfile],
        total_page_count: int,
    ) -> None:
        if not page_profiles:
            base.pdf_type = "unknown"
            base.processing_strategy = "manual_review"
            base.complexity_level = "unknown"
            base.need_ocr = None
            base.warnings.append("No page profile was generated.")
            return

        sample_count = len(page_profiles)

        text_pages = [p for p in page_profiles if p.has_text]
        image_pages = [p for p in page_profiles if p.has_image]
        scanned_pages = [p for p in page_profiles if p.likely_scanned]
        ocr_pages = [p for p in page_profiles if p.need_ocr]
        table_pages = [p for p in page_profiles if p.has_table_candidate]

        text_ratio = len(text_pages) / sample_count
        image_ratio = len(image_pages) / sample_count
        scanned_ratio = len(scanned_pages) / sample_count
        table_ratio = len(table_pages) / sample_count
        ocr_ratio = len(ocr_pages) / sample_count

        avg_text_chars = self._safe_mean([p.text_char_count for p in page_profiles])
        avg_blocks = self._safe_mean([p.text_block_count for p in page_profiles])
        avg_images = self._safe_mean([p.image_count for p in page_profiles])
        avg_drawings = self._safe_mean([p.drawing_count for p in page_profiles])
        avg_image_area_ratio = self._safe_mean(
            [p.image_area_ratio for p in page_profiles]
        )

        base.has_text_layer = text_ratio > 0.2
        base.has_image_layer = image_ratio > 0.2
        base.need_ocr = ocr_ratio > 0.2

        base.pdf_type = self._classify_pdf_type(
            text_ratio=text_ratio,
            image_ratio=image_ratio,
            scanned_ratio=scanned_ratio,
        )

        base.processing_strategy = self._recommend_processing_strategy(
            pdf_type=base.pdf_type,
            table_ratio=table_ratio,
            need_ocr=bool(base.need_ocr),
        )

        base.complexity_level = self._classify_document_complexity(
            table_ratio=table_ratio,
            image_ratio=image_ratio,
            avg_blocks=avg_blocks,
        )

        base.ocr_page_numbers = [p.page_number for p in ocr_pages]
        base.table_candidate_pages = [p.page_number for p in table_pages]
        base.image_pages = [p.page_number for p in image_pages]

        base.language_hint = self._detect_language_hint_from_profiles(page_profiles)

        base.stats = {
            "total_page_count": total_page_count,
            "sampled_page_count": sample_count,
            "text_page_ratio": round(text_ratio, 4),
            "image_page_ratio": round(image_ratio, 4),
            "scanned_page_ratio": round(scanned_ratio, 4),
            "ocr_page_ratio": round(ocr_ratio, 4),
            "table_candidate_page_ratio": round(table_ratio, 4),
            "avg_text_chars_per_page": round(avg_text_chars, 2),
            "avg_text_blocks_per_page": round(avg_blocks, 2),
            "avg_images_per_page": round(avg_images, 2),
            "avg_drawings_per_page": round(avg_drawings, 2),
            "avg_image_area_ratio": round(avg_image_area_ratio, 4),
        }

    # --------------------------------------------------------
    # DOCX profiling
    # --------------------------------------------------------

    def _profile_docx(
        self,
        document_path: Union[str, Path],
        base: DocumentProfile,
        loaded_doc: Any,
    ) -> DocumentProfile:
        if DocxDocument is None:
            raise MissingDependencyError(
                "python-docx is not installed. Run: pip install python-docx"
            )

        doc = DocxDocument(str(document_path))

        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        paragraph_count = len(paragraphs)
        table_count = len(doc.tables)

        total_chars = sum(len(p) for p in paragraphs)
        total_words = sum(len(p.split()) for p in paragraphs)

        heading_count = 0
        for p in doc.paragraphs:
            style_name = p.style.name.lower() if p.style and p.style.name else ""
            if "heading" in style_name or "title" in style_name:
                heading_count += 1

        base.has_text_layer = total_chars > 0
        base.has_image_layer = None
        base.need_ocr = False
        base.pdf_type = None
        base.language_hint = self._detect_language_hint("\n".join(paragraphs[:20]))

        base.processing_strategy = "docx_structure_extraction"
        base.complexity_level = self._classify_docx_complexity(
            paragraph_count=paragraph_count,
            table_count=table_count,
            heading_count=heading_count,
        )

        base.table_candidate_pages = []
        base.ocr_page_numbers = []

        base.stats = {
            "paragraph_count": paragraph_count,
            "table_count": table_count,
            "heading_count": heading_count,
            "total_chars": total_chars,
            "total_words": total_words,
            "avg_chars_per_paragraph": round(
                total_chars / paragraph_count, 2
            ) if paragraph_count else 0,
        }

        base.metadata.update(
            {
                "docx_profile": {
                    "paragraph_count": paragraph_count,
                    "table_count": table_count,
                    "heading_count": heading_count,
                }
            }
        )

        return base

    # --------------------------------------------------------
    # Image profiling
    # --------------------------------------------------------

    def _profile_image(
        self,
        document_path: Union[str, Path],
        base: DocumentProfile,
        loaded_doc: Any,
    ) -> DocumentProfile:
        if Image is None:
            raise MissingDependencyError(
                "Pillow is not installed. Run: pip install pillow"
            )

        with Image.open(str(document_path)) as img:
            width, height = img.size

            base.has_text_layer = False
            base.has_image_layer = True
            base.need_ocr = True
            base.pdf_type = None
            base.language_hint = None
            base.processing_strategy = "image_ocr"
            base.complexity_level = "medium"

            base.ocr_page_numbers = [1]

            base.stats = {
                "width": width,
                "height": height,
                "mode": img.mode,
                "format": img.format,
                "n_frames": getattr(img, "n_frames", 1),
                "pixel_count": width * height,
            }

            base.pages = [
                {
                    "page_number": 1,
                    "width": width,
                    "height": height,
                    "has_text": False,
                    "has_image": True,
                    "need_ocr": True,
                    "page_type": "image",
                }
            ]

        return base

    # --------------------------------------------------------
    # Text profiling
    # --------------------------------------------------------

    def _profile_text(
        self,
        document_path: Union[str, Path],
        base: DocumentProfile,
        loaded_doc: Any,
    ) -> DocumentProfile:
        text_preview = ""
        if loaded_doc.metadata:
            text_preview = loaded_doc.metadata.get("text_preview", "")

        base.has_text_layer = True
        base.has_image_layer = False
        base.need_ocr = False
        base.pdf_type = None
        base.language_hint = self._detect_language_hint(text_preview)
        base.processing_strategy = "plain_text_extraction"
        base.complexity_level = "low"

        base.stats = {
            "line_count": loaded_doc.metadata.get("line_count")
            if loaded_doc.metadata else None,
            "char_count": loaded_doc.metadata.get("char_count")
            if loaded_doc.metadata else None,
        }

        base.pages = [
            {
                "page_number": 1,
                "has_text": True,
                "has_image": False,
                "need_ocr": False,
                "page_type": "text",
            }
        ]

        return base

    # --------------------------------------------------------
    # PDF helper methods
    # --------------------------------------------------------

    def _select_page_indices(
        self,
        page_count: int,
        profile_all_pages: bool = False,
    ) -> List[int]:
        if page_count <= 0:
            return []

        if profile_all_pages:
            return list(range(page_count))

        sample_size = min(self.config.sample_pages, page_count)

        if sample_size >= page_count:
            return list(range(page_count))

        indices = set()

        # Always include first pages
        for i in range(min(3, page_count)):
            indices.add(i)

        # Include middle pages
        middle = page_count // 2
        for i in range(max(0, middle - 1), min(page_count, middle + 2)):
            indices.add(i)

        # Include last pages
        for i in range(max(0, page_count - 3), page_count):
            indices.add(i)

        # Fill remaining positions evenly
        if len(indices) < sample_size:
            step = max(page_count // sample_size, 1)
            for i in range(0, page_count, step):
                indices.add(i)
                if len(indices) >= sample_size:
                    break

        return sorted(list(indices))[:sample_size]

    def _count_text_blocks(self, text_dict: Dict[str, Any]) -> int:
        blocks = text_dict.get("blocks", []) if text_dict else []

        count = 0
        for block in blocks:
            if block.get("type") == 0:
                count += 1

        return count

    def _estimate_image_area_ratio(self, page: Any, page_area: float) -> float:
        """
        Estimate image coverage ratio on page.

        PyMuPDF image metadata does not always directly provide displayed bbox.
        This method uses page.get_text("dict") image blocks when available.
        """

        try:
            text_dict = page.get_text("dict") or {}
            blocks = text_dict.get("blocks", [])

            image_area = 0.0

            for block in blocks:
                if block.get("type") == 1:
                    bbox = block.get("bbox")
                    if bbox and len(bbox) == 4:
                        x0, y0, x1, y1 = bbox
                        image_area += max((x1 - x0) * (y1 - y0), 0)

            return min(image_area / page_area, 1.0)

        except Exception:
            return 0.0

    def _estimate_table_candidates(
        self,
        text: str,
        drawings: List[Any],
    ) -> int:
        """
        Lightweight table candidate heuristic.

        It does not replace TableBoundaryDetector.
        It only estimates whether a page likely contains tables.
        """

        candidate_score = 0

        lines = [line.strip() for line in text.splitlines() if line.strip()]

        aligned_like_lines = 0
        numeric_like_lines = 0

        for line in lines:
            # Many spaces or tab-like separation
            if "\t" in line or "  " in line:
                aligned_like_lines += 1

            # Numeric-heavy lines often appear in tables
            digit_count = sum(ch.isdigit() for ch in line)
            if len(line) > 0 and digit_count / max(len(line), 1) > 0.25:
                numeric_like_lines += 1

        if aligned_like_lines >= self.config.table_min_aligned_lines:
            candidate_score += 1

        if numeric_like_lines >= self.config.table_min_numeric_lines:
            candidate_score += 1

        # Many drawings/lines usually indicate ruled tables
        if drawings and len(drawings) >= 10:
            candidate_score += 1

        return candidate_score

    def _classify_page_type(
        self,
        has_text: bool,
        has_image: bool,
        has_table_candidate: bool,
        likely_scanned: bool,
    ) -> str:
        if likely_scanned:
            return "scanned_page"

        if has_table_candidate and has_text:
            return "text_with_table"

        if has_text and has_image:
            return "text_with_image"

        if has_text:
            return "text_page"

        if has_image:
            return "image_page"

        return "blank_or_unknown"

    def _classify_page_complexity(
        self,
        text_block_count: int,
        image_count: int,
        drawing_count: int,
        table_candidate_count: int,
    ) -> str:
        score = 0

        if text_block_count > 80:
            score += 2
        elif text_block_count > 40:
            score += 1

        if image_count > 5:
            score += 2
        elif image_count > 0:
            score += 1

        if drawing_count > 30:
            score += 2
        elif drawing_count > 10:
            score += 1

        if table_candidate_count > 1:
            score += 2
        elif table_candidate_count == 1:
            score += 1

        if score >= 5:
            return "high"

        if score >= 2:
            return "medium"

        return "low"

    def _classify_pdf_type(
        self,
        text_ratio: float,
        image_ratio: float,
        scanned_ratio: float,
    ) -> str:
        if text_ratio >= 0.8 and scanned_ratio < 0.2:
            return "digital_pdf"

        if scanned_ratio >= 0.7:
            return "scanned_pdf"

        if text_ratio > 0.2 and image_ratio > 0.2:
            return "hybrid_pdf"

        if text_ratio <= 0.2 and image_ratio <= 0.2:
            return "blank_or_unknown_pdf"

        return "mixed_pdf"

    def _recommend_processing_strategy(
        self,
        pdf_type: str,
        table_ratio: float,
        need_ocr: bool,
    ) -> str:
        if pdf_type == "digital_pdf":
            if table_ratio >= 0.2:
                return "text_extraction_plus_table_detection"
            return "text_extraction"

        if pdf_type == "scanned_pdf":
            return "full_page_ocr"

        if pdf_type in {"hybrid_pdf", "mixed_pdf"}:
            if need_ocr:
                return "hybrid_text_extraction_region_ocr"
            return "hybrid_text_extraction"

        return "manual_review"

    def _classify_document_complexity(
        self,
        table_ratio: float,
        image_ratio: float,
        avg_blocks: float,
    ) -> str:
        if (
            table_ratio >= self.config.high_complexity_table_ratio
            or image_ratio >= self.config.high_complexity_image_ratio
            or avg_blocks >= self.config.high_complexity_avg_blocks
        ):
            return "high"

        if table_ratio > 0 or image_ratio > 0.2 or avg_blocks > 40:
            return "medium"

        return "low"

    # --------------------------------------------------------
    # DOCX helper methods
    # --------------------------------------------------------

    def _classify_docx_complexity(
        self,
        paragraph_count: int,
        table_count: int,
        heading_count: int,
    ) -> str:
        if table_count >= 10 or paragraph_count >= 300:
            return "high"

        if table_count > 0 or paragraph_count >= 80 or heading_count >= 10:
            return "medium"

        return "low"

    # --------------------------------------------------------
    # Language hint
    # --------------------------------------------------------

    def _detect_language_hint_from_profiles(
        self,
        page_profiles: List[PageProfile],
    ) -> Optional[str]:
        # PageProfile does not keep full text to avoid bloating JSON.
        # Language detection will be handled better in TextExtractor.
        # Here we return None unless extended later.
        return None

    def _detect_language_hint(self, text: str) -> Optional[str]:
        """
        Very lightweight language hint.

        This is intentionally simple and dependency-free.
        For production, replace with fastText/langdetect/cld3.
        """

        if not text:
            return None

        vietnamese_chars = set(
            "ăâđêôơưáàảãạấầẩẫậắằẳẵặ"
            "éèẻẽẹếềểễệíìỉĩị"
            "óòỏõọốồổỗộớờởỡợ"
            "úùủũụứừửữựýỳỷỹỵ"
            "ĂÂĐÊÔƠƯÁÀẢÃẠẤẦẨẪẬẮẰẲẴẶ"
            "ÉÈẺẼẸẾỀỂỄỆÍÌỈĨỊ"
            "ÓÒỎÕỌỐỒỔỖỘỚỜỞỠỢ"
            "ÚÙỦŨỤỨỪỬỮỰÝỲỶỸỴ"
        )

        vn_count = sum(1 for ch in text if ch in vietnamese_chars)

        if vn_count >= 5:
            return "vi"

        ascii_letters = sum(1 for ch in text if ch.isascii() and ch.isalpha())
        if ascii_letters > 20:
            return "en_or_latin"

        return "unknown"

    # --------------------------------------------------------
    # Utility
    # --------------------------------------------------------

    def _safe_mean(self, values: List[Union[int, float]]) -> float:
        if not values:
            return 0.0

        return float(statistics.mean(values))


# ============================================================
# CLI-like helper for Colab
# ============================================================

def profile_document(
    document_path: Union[str, Path],
    profile_all_pages: bool = False,
) -> Dict[str, Any]:
    """
    Convenience function for notebook usage.
    """

    profiler = DocumentProfiler()
    profile = profiler.process(
        document_path=document_path,
        profile_all_pages=profile_all_pages,
    )

    return profile.to_dict()
