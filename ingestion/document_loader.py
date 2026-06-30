from __future__ import annotations

import os
import json
import uuid
import hashlib
import mimetypes
from pathlib import Path
from dataclasses import dataclass, asdict, field
from datetime import datetime, timezone
from typing import Any, Dict, Iterator, List, Optional, Union
from contextlib import contextmanager


# ============================================================
# Optional dependencies
# ============================================================

try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

try:
    from docx import Document as DocxDocument
except Exception:
    DocxDocument = None

try:
    from PIL import Image
except Exception:
    Image = None


# ============================================================
# Exceptions
# ============================================================

class DocumentLoaderError(Exception):
    """Base exception for document loader."""


class UnsupportedDocumentTypeError(DocumentLoaderError):
    """Raised when file type is not supported."""


class MissingDependencyError(DocumentLoaderError):
    """Raised when required dependency is missing."""


class FileValidationError(DocumentLoaderError):
    """Raised when file validation fails."""


# ============================================================
# Data classes
# ============================================================

@dataclass
class DocumentLoaderConfig:
    """
    Configuration for DocumentLoader.
    """

    max_file_size_mb: Optional[int] = None
    allow_encrypted_pdf: bool = False
    sample_pages_for_pdf_profile: int = 5
    read_text_preview_chars: int = 3000


@dataclass
class LoadedDocument:
    """
    JSON-safe document loading result.
    """

    document_id: str
    source_path: str
    file_name: str
    file_extension: str
    document_type: str
    mime_type: str
    file_size_bytes: int
    sha256: str

    page_count: Optional[int] = None
    is_encrypted: bool = False
    need_ocr: Optional[bool] = None
    has_text_layer: Optional[bool] = None
    has_image_layer: Optional[bool] = None

    metadata: Dict[str, Any] = field(default_factory=dict)
    warnings: List[str] = field(default_factory=list)

    loaded_at_utc: str = field(
        default_factory=lambda: datetime.now(timezone.utc).isoformat()
    )

    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)

    def to_json(self, ensure_ascii: bool = False, indent: int = 2) -> str:
        return json.dumps(self.to_dict(), ensure_ascii=ensure_ascii, indent=indent)


# ============================================================
# Document Loader
# ============================================================

class DocumentLoader:
    """
    Load and inspect input documents.

    This class should not perform deep extraction.
    It only validates file, detects type, and returns basic metadata.

    Deep extraction should be done by:
    - TextExtractor
    - ImageExtractor
    - TableExtractor
    - DrawingExtractor
    - PageExtractionPipeline
    """

    IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}
    TEXT_EXTENSIONS = {".txt", ".md"}
    DOCX_EXTENSIONS = {".docx"}
    PDF_EXTENSIONS = {".pdf"}

    SUPPORTED_EXTENSIONS = (
        IMAGE_EXTENSIONS
        | TEXT_EXTENSIONS
        | DOCX_EXTENSIONS
        | PDF_EXTENSIONS
    )

    def __init__(self, config: Optional[DocumentLoaderConfig] = None):
        self.config = config or DocumentLoaderConfig()

    # --------------------------------------------------------
    # Public API
    # --------------------------------------------------------

    def load(
        self,
        document_path: Union[str, Path],
        document_id: Optional[str] = None
    ) -> LoadedDocument:
        """
        Load a document and return JSON-safe metadata.

        Parameters
        ----------
        document_path:
            Path to document file.

        document_id:
            Optional custom document id.

        Returns
        -------
        LoadedDocument
        """

        path = self._resolve_path(document_path)
        self._validate_file(path)

        file_extension = path.suffix.lower()
        document_type = self._detect_document_type(path)
        mime_type = self._detect_mime_type(path)
        file_size_bytes = path.stat().st_size
        sha256 = self._compute_sha256(path)

        base = LoadedDocument(
            document_id=document_id or self._generate_document_id(path),
            source_path=str(path),
            file_name=path.name,
            file_extension=file_extension,
            document_type=document_type,
            mime_type=mime_type,
            file_size_bytes=file_size_bytes,
            sha256=sha256,
        )

        if document_type == "pdf":
            return self._load_pdf(path, base)

        if document_type == "docx":
            return self._load_docx(path, base)

        if document_type == "image":
            return self._load_image(path, base)

        if document_type == "text":
            return self._load_text(path, base)

        raise UnsupportedDocumentTypeError(
            f"Unsupported document type: {document_type}"
        )

    @contextmanager
    def open_pdf(self, document_path: Union[str, Path]):
        """
        Context manager to open PDF safely.

        Usage
        -----
        with loader.open_pdf("sample.pdf") as pdf:
            page = pdf.load_page(0)
        """

        if fitz is None:
            raise MissingDependencyError(
                "PyMuPDF is not installed. Run: pip install pymupdf"
            )

        path = self._resolve_path(document_path)
        self._validate_file(path)

        if path.suffix.lower() != ".pdf":
            raise UnsupportedDocumentTypeError("open_pdf only supports PDF files.")

        doc = fitz.open(str(path))

        try:
            if doc.is_encrypted and not self.config.allow_encrypted_pdf:
                raise DocumentLoaderError(
                    "PDF is encrypted. Set allow_encrypted_pdf=True if needed."
                )

            yield doc

        finally:
            doc.close()

    def iter_pdf_pages(
        self,
        document_path: Union[str, Path],
        start_page: int = 1,
        end_page: Optional[int] = None,
        max_pages: Optional[int] = None
    ) -> Iterator[Dict[str, Any]]:
        """
        Iterate PDF pages.

        Page number is 1-based.

        Yields
        ------
        {
            "page_number": int,
            "page_index": int,
            "width": float,
            "height": float,
            "rotation": int,
            "page": fitz.Page
        }
        """

        if fitz is None:
            raise MissingDependencyError(
                "PyMuPDF is not installed. Run: pip install pymupdf"
            )

        path = self._resolve_path(document_path)

        with fitz.open(str(path)) as doc:
            if doc.is_encrypted and not self.config.allow_encrypted_pdf:
                raise DocumentLoaderError(
                    "PDF is encrypted. Set allow_encrypted_pdf=True if needed."
                )

            total_pages = doc.page_count

            start_index = max(start_page - 1, 0)
            end_index = total_pages if end_page is None else min(end_page, total_pages)

            yielded = 0

            for page_index in range(start_index, end_index):
                if max_pages is not None and yielded >= max_pages:
                    break

                page = doc.load_page(page_index)
                rect = page.rect

                yield {
                    "page_number": page_index + 1,
                    "page_index": page_index,
                    "width": float(rect.width),
                    "height": float(rect.height),
                    "rotation": int(page.rotation),
                    "page": page,
                }

                yielded += 1

    # --------------------------------------------------------
    # Loaders by type
    # --------------------------------------------------------

    def _load_pdf(self, path: Path, base: LoadedDocument) -> LoadedDocument:
        if fitz is None:
            raise MissingDependencyError(
                "PyMuPDF is not installed. Run: pip install pymupdf"
            )

        warnings: List[str] = []

        with fitz.open(str(path)) as doc:
            base.page_count = doc.page_count
            base.is_encrypted = bool(doc.is_encrypted)

            if doc.is_encrypted:
                if not self.config.allow_encrypted_pdf:
                    raise DocumentLoaderError(
                        "PDF is encrypted. Set allow_encrypted_pdf=True if needed."
                    )

                warnings.append("PDF is encrypted.")

            pdf_metadata = doc.metadata or {}

            sample_pages = min(
                self.config.sample_pages_for_pdf_profile,
                doc.page_count
            )

            text_chars = 0
            image_count = 0

            for page_index in range(sample_pages):
                page = doc.load_page(page_index)

                try:
                    page_text = page.get_text("text") or ""
                    text_chars += len(page_text.strip())
                except Exception:
                    warnings.append(
                        f"Could not extract text from page {page_index + 1}."
                    )

                try:
                    image_count += len(page.get_images(full=True))
                except Exception:
                    warnings.append(
                        f"Could not inspect images from page {page_index + 1}."
                    )

            base.has_text_layer = text_chars > 20
            base.has_image_layer = image_count > 0
            base.need_ocr = not base.has_text_layer

            base.metadata = {
                "pdf_metadata": pdf_metadata,
                "sampled_pages": sample_pages,
                "sample_text_chars": text_chars,
                "sample_image_count": image_count,
                "producer": pdf_metadata.get("producer"),
                "creator": pdf_metadata.get("creator"),
                "author": pdf_metadata.get("author"),
                "title": pdf_metadata.get("title"),
                "subject": pdf_metadata.get("subject"),
                "keywords": pdf_metadata.get("keywords"),
            }

            base.warnings = warnings

        return base

    def _load_docx(self, path: Path, base: LoadedDocument) -> LoadedDocument:
        if DocxDocument is None:
            raise MissingDependencyError(
                "python-docx is not installed. Run: pip install python-docx"
            )

        doc = DocxDocument(str(path))

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text_preview = "\n".join(paragraphs)

        base.page_count = None
        base.is_encrypted = False
        base.has_text_layer = True
        base.has_image_layer = None
        base.need_ocr = False

        core = doc.core_properties

        base.metadata = {
            "paragraph_count": len(doc.paragraphs),
            "non_empty_paragraph_count": len(paragraphs),
            "table_count": len(doc.tables),
            "text_preview": text_preview[: self.config.read_text_preview_chars],
            "core_properties": {
                "author": core.author,
                "title": core.title,
                "subject": core.subject,
                "keywords": core.keywords,
                "comments": core.comments,
                "category": core.category,
                "created": str(core.created) if core.created else None,
                "modified": str(core.modified) if core.modified else None,
                "last_modified_by": core.last_modified_by,
            },
        }

        return base

    def _load_image(self, path: Path, base: LoadedDocument) -> LoadedDocument:
        if Image is None:
            raise MissingDependencyError(
                "Pillow is not installed. Run: pip install pillow"
            )

        with Image.open(str(path)) as img:
            width, height = img.size
            n_frames = getattr(img, "n_frames", 1)

            base.page_count = n_frames
            base.is_encrypted = False
            base.has_text_layer = False
            base.has_image_layer = True
            base.need_ocr = True

            base.metadata = {
                "image_format": img.format,
                "width": width,
                "height": height,
                "mode": img.mode,
                "n_frames": n_frames,
                "dpi": img.info.get("dpi"),
            }

        return base

    def _load_text(self, path: Path, base: LoadedDocument) -> LoadedDocument:
        text = self._read_text_file(path)

        base.page_count = 1
        base.is_encrypted = False
        base.has_text_layer = True
        base.has_image_layer = False
        base.need_ocr = False

        base.metadata = {
            "line_count": len(text.splitlines()),
            "char_count": len(text),
            "text_preview": text[: self.config.read_text_preview_chars],
        }

        return base

    # --------------------------------------------------------
    # Helpers
    # --------------------------------------------------------

    def _resolve_path(self, document_path: Union[str, Path]) -> Path:
        path = Path(document_path).expanduser()

        if not path.is_absolute():
            path = Path.cwd() / path

        return path.resolve()

    def _validate_file(self, path: Path) -> None:
        if not path.exists():
            raise FileValidationError(f"File does not exist: {path}")

        if not path.is_file():
            raise FileValidationError(f"Path is not a file: {path}")

        if path.suffix.lower() not in self.SUPPORTED_EXTENSIONS:
            raise UnsupportedDocumentTypeError(
                f"Unsupported file extension: {path.suffix}. "
                f"Supported: {sorted(self.SUPPORTED_EXTENSIONS)}"
            )

        if self.config.max_file_size_mb is not None:
            max_bytes = self.config.max_file_size_mb * 1024 * 1024
            file_size = path.stat().st_size

            if file_size > max_bytes:
                raise FileValidationError(
                    f"File is too large: {file_size} bytes. "
                    f"Limit: {max_bytes} bytes."
                )

    def _detect_document_type(self, path: Path) -> str:
        ext = path.suffix.lower()

        if ext in self.PDF_EXTENSIONS:
            return "pdf"

        if ext in self.DOCX_EXTENSIONS:
            return "docx"

        if ext in self.IMAGE_EXTENSIONS:
            return "image"

        if ext in self.TEXT_EXTENSIONS:
            return "text"

        raise UnsupportedDocumentTypeError(f"Unsupported extension: {ext}")

    def _detect_mime_type(self, path: Path) -> str:
        mime_type, _ = mimetypes.guess_type(str(path))
        return mime_type or "application/octet-stream"

    def _compute_sha256(self, path: Path, chunk_size: int = 1024 * 1024) -> str:
        sha = hashlib.sha256()

        with open(path, "rb") as f:
            while True:
                chunk = f.read(chunk_size)
                if not chunk:
                    break
                sha.update(chunk)

        return sha.hexdigest()

    def _generate_document_id(self, path: Path) -> str:
        raw = f"{path.name}-{path.stat().st_size}-{path.stat().st_mtime}"
        return "doc_" + uuid.uuid5(uuid.NAMESPACE_URL, raw).hex[:16]

    def _read_text_file(self, path: Path) -> str:
        encodings = ["utf-8", "utf-8-sig", "cp1258", "latin-1"]

        for enc in encodings:
            try:
                with open(path, "r", encoding=enc) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue

        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
